from __future__ import annotations

import io
import logging
import uuid
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import BASE_DIR
from app.models.chapter import Chapter
from app.models.project import Project
from app.models.volume import Volume

logger = logging.getLogger(__name__)

EXPORT_DIR = BASE_DIR / "data" / "exports"


async def export_project(
    session: AsyncSession,
    project_id: uuid.UUID,
    fmt: str = "txt",
    scope: str = "all",
    volume_id: uuid.UUID | None = None,
    chapter_id: uuid.UUID | None = None,
) -> tuple[Path, str]:
    """Export project content to file. Returns (file_path, filename)."""
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    project = await _load_project(session, project_id)
    if not project:
        raise ValueError("项目不存在")

    chapters = await _get_export_chapters(session, project_id, scope, volume_id, chapter_id)
    if not chapters:
        raise ValueError("没有可导出的章节内容")

    title = project.name or "未命名小说"

    if fmt == "txt":
        return _export_txt(title, chapters)
    elif fmt == "docx":
        return _export_docx(title, chapters)
    elif fmt == "epub":
        return _export_epub(title, chapters, project)
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")


def _export_txt(title: str, chapters: list[Chapter]) -> tuple[Path, str]:
    filename = f"{title}.txt"
    filepath = EXPORT_DIR / f"{uuid.uuid4()}.txt"

    lines = [title, "=" * len(title.encode("utf-8")), ""]
    for ch in chapters:
        lines.append(f"\n{ch.title}\n")
        lines.append(ch.content or "（暂无内容）")
        lines.append("")

    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filepath, filename


def _export_docx(title: str, chapters: list[Chapter]) -> tuple[Path, str]:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=0)

    for ch in chapters:
        doc.add_heading(ch.title, level=1)
        content = ch.content or "（暂无内容）"
        for para_text in content.split("\n"):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(12)

    filename = f"{title}.docx"
    filepath = EXPORT_DIR / f"{uuid.uuid4()}.docx"
    doc.save(str(filepath))
    return filepath, filename


def _export_epub(
    title: str, chapters: list[Chapter], project: Project
) -> tuple[Path, str]:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language("zh")
    book.add_author(project.ai_provider or "千卷 AI")

    spine = ["nav"]
    toc = []

    for i, ch in enumerate(chapters):
        content = ch.content or "（暂无内容）"
        html_content = "".join(f"<p>{line}</p>" for line in content.split("\n") if line.strip())

        epub_ch = epub.EpubHtml(
            title=ch.title,
            file_name=f"chapter_{i+1}.xhtml",
            lang="zh",
        )
        epub_ch.content = f"<h1>{ch.title}</h1>{html_content}"
        book.add_item(epub_ch)
        spine.append(epub_ch)
        toc.append(epub_ch)

    book.toc = toc
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    filename = f"{title}.epub"
    filepath = EXPORT_DIR / f"{uuid.uuid4()}.epub"
    epub.write_epub(str(filepath), book)
    return filepath, filename


async def _load_project(session: AsyncSession, project_id: uuid.UUID) -> Project | None:
    return await session.get(Project, project_id)


async def _get_export_chapters(
    session: AsyncSession,
    project_id: uuid.UUID,
    scope: str,
    volume_id: uuid.UUID | None,
    chapter_id: uuid.UUID | None,
) -> list[Chapter]:
    if scope == "chapter" and chapter_id:
        ch = await session.get(Chapter, chapter_id)
        return [ch] if ch and ch.content else []

    if scope == "volume" and volume_id:
        stmt = (
            select(Chapter)
            .where(Chapter.volume_id == volume_id, Chapter.content.isnot(None))
            .order_by(Chapter.sort_order)
        )
    else:
        stmt = (
            select(Chapter)
            .where(Chapter.project_id == project_id, Chapter.content.isnot(None))
            .order_by(Chapter.sort_order)
        )

    result = await session.execute(stmt)
    return list(result.scalars().all())
